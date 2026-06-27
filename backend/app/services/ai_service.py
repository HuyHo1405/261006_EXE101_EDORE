"""
ai_service.py — Business logic cho AI pipeline.

Tách ra từ routes/ai.py và routes/stream.py.
Controller chỉ parse request và gọi vào đây.
"""

import io
import re
import json
import time
from typing import Optional, Generator
from concurrent.futures import ThreadPoolExecutor, as_completed

try:
    import pypdf
    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

try:
    import docx as python_docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from app.services.openrouter_service import OpenRouterService
from app.services.chunking_service import get_all_section_contexts, semantic_chunk, extract_outline
from app.models.template_store import TemplateStore


# ─── SSE Event Formatters ─────────────────────────────────────────────────────

def sse_event(event_type: str, data: dict) -> str:
    """Format một SSE event string theo chuẩn EventSource."""
    return f"event: {event_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def sse_progress(message: str, step: int, total_steps: int) -> str:
    return sse_event("progress", {
        "message": message, "step": step,
        "total_steps": total_steps, "timestamp": time.time(),
    })


def sse_section(node_data: dict, section_index: int) -> str:
    return sse_event("section", {
        "index": section_index, "node": node_data, "timestamp": time.time(),
    })


def sse_done(summary: dict) -> str:
    return sse_event("done", {**summary, "timestamp": time.time()})


def sse_error(message: str, details: str = None) -> str:
    return sse_event("error", {
        "message": message, "details": details, "timestamp": time.time(),
    })


def sse_raw_data(raw_text: str) -> str:
    return sse_event("raw_data", {"raw_text": raw_text, "timestamp": time.time()})


def sse_chunks(chunks: list) -> str:
    return sse_event("chunks", {"chunks": chunks, "timestamp": time.time()})


def sse_content_summary(summary: str) -> str:
    return sse_event("content_summary", {"summary": summary, "timestamp": time.time()})


# ─── Default Template Data ────────────────────────────────────────────────────

DEFAULT_NODES = [
    {"node_type": "Khởi động", "goal": "Kích hoạt kiến thức nền và tạo hứng thú cho học sinh"},
    {"node_type": "Hình thành kiến thức", "goal": "Giới thiệu, giải thích và xây dựng các khái niệm, kiến thức mới cốt lõi"},
    {"node_type": "Luyện tập", "goal": "Học sinh thực hành, làm các bài tập củng cố ngay tại lớp"},
]

DEFAULT_RAG_ACTIVITIES = [
    "1. Trò chơi Kahoot câu hỏi trắc nghiệm ôn tập nhanh (5-7 phút)",
    "2. Thảo luận nhóm tranh biện (Think-Pair-Share)",
    "3. Sơ đồ tư duy tiếp sức theo nhóm (Mind Map Relay)",
    "4. Nhập vai xử lý tình huống thực tế (Role Play)",
    "5. Thí nghiệm hoặc mô phỏng thực hành có hướng dẫn",
]

DEFAULT_SYSTEM_TEMPLATE = (
    "Khung bài học 3 node:\n"
    "Node 1 — node_type: 'Khởi động' | Mục tiêu: Kích hoạt kiến thức nền và tạo hứng thú cho học sinh\n"
    "Node 2 — node_type: 'Hình thành kiến thức' | Mục tiêu: Giới thiệu, giải thích và xây dựng các khái niệm, kiến thức mới cốt lõi\n"
    "Node 3 — node_type: 'Luyện tập' | Mục tiêu: Học sinh thực hành, làm các bài tập củng cố ngay tại lớp\n\n"
    "QUY TẮC: Trường 'title' phải là tiêu đề ngắn gọn mô tả NỘI DUNG CỤ THỂ của node này trong bài học (ĐỪNG thêm vào 'node_type')."
)

DEFAULT_NODE_NAMES = ["Khởi động (Warm-up)", "Hình thành kiến thức (Core Theory)", "Luyện tập (Practice)"]



# ─── AIService ────────────────────────────────────────────────────────────────

class AIService:

    # ── Text Extraction ───────────────────────────────────────────────────────

    @staticmethod
    def extract_text_from_file(file) -> tuple:
        """
        Extract text từ PDF / DOCX / TXT / MD.
        Returns (text: str | None, error: str | None)
        """
        filename = file.filename.lower()
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

        try:
            file.seek(0, 2)
            size = file.tell()
            if size > MAX_FILE_SIZE:
                return None, "Kích thước file vượt quá giới hạn cho phép (tối đa 50MB)."
        except Exception:
            pass
        finally:
            try:
                file.seek(0)
            except Exception:
                pass

        try:
            if filename.endswith('.pdf'):
                # Reset pointer to 0 just in case
                try:
                    file.seek(0)
                except Exception:
                    pass
                import fitz  # PyMuPDF
                pdf_bytes = file.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    t = page.get_text()
                    if t:
                        text += t + "\n"

                # Khắc phục lỗi newline phân mảnh — dùng heuristic để phân biệt 2 loại PDF:
                #   Loại A (fragmented): mỗi ký tự / từ nằm trên 1 dòng riêng → cần join
                #   Loại B (structured): PDF chuẩn có heading BÀI/MỤC trên dòng riêng → giữ nguyên
                if text.strip():
                    non_empty_lines = [l for l in text.split('\n') if l.strip()]
                    avg_line_len = (
                        sum(len(l.strip()) for l in non_empty_lines) / len(non_empty_lines)
                        if non_empty_lines else 0
                    )

                    if avg_line_len < 4.0:  # Hạ ngưỡng để tránh nhận nhầm PDF chuẩn
                        # Loại A: char-per-line / word-per-line fragmentation
                        # Gom hết thành 1 dòng rồi để sgk_parser._preprocess tách lại theo heading
                        text = re.sub(r'\n[ \t]*\n', '\n\n', text)   # giữ double-newline
                        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)  # single \n → space
                        text = re.sub(r'[ \t]{2,}', ' ', text)        # collapse spaces
                    else:
                        # Loại B: PDF có cấu trúc — CHỈ dọn blank lines & trailing spaces
                        text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)  # trailing spaces
                        text = re.sub(r'\n{3,}', '\n\n', text)                  # max 2 blank lines

                # Dự phòng OCR: Nếu text rỗng hoặc quá ngắn (< 50 ký tự), tiến hành chạy OCR
                if not text.strip() or len(text.strip()) < 50:
                    try:
                        import pypdfium2 as pdfium
                        import easyocr
                        import numpy as np

                        # Khởi tạo EasyOCR Reader cho tiếng Việt ('vi') và tiếng Anh ('en')
                        # Lần đầu tiên chạy sẽ tự động tải file model OCR (~100MB) về máy
                        ocr_reader = easyocr.Reader(['vi', 'en'])

                        pdf_bytes.seek(0)
                        doc = pdfium.PdfDocument(pdf_bytes)
                        total_pages = len(doc)
                        ocr_pages = []

                        print(f"[OCR] Phát hiện PDF dạng ảnh quét. Bắt đầu chạy OCR cho {total_pages} trang...")

                        for i, page in enumerate(doc):
                            print(f"[OCR] Đang xử lý trang {i + 1}/{total_pages}...")
                            # Render trang thành ảnh với độ phân giải vừa phải (scale=1.2) để chạy nhanh hơn trên CPU
                            bitmap = page.render(scale=1.2)
                            pil_img = bitmap.to_pil()
                            img_np = np.array(pil_img)
                            
                            # Nhận diện chữ
                            results = ocr_reader.readtext(img_np, detail=0)
                            page_text = " ".join(results).strip()
                            if page_text:
                                ocr_pages.append(page_text)

                        if ocr_pages:
                            text = "\n\n".join(ocr_pages)
                            print(f"[OCR] Hoàn thành OCR cho {total_pages} trang thành công!")
                    except Exception as ocr_err:
                        print(f"[OCR] Lỗi trong quá trình chạy OCR: {str(ocr_err)}")
                        return None, f"Không thể trích xuất chữ thường, và chạy OCR dự phòng gặp lỗi: {str(ocr_err)}"


            elif filename.endswith(('.txt', '.md')):
                text = file.read().decode('utf-8', errors='ignore')

            elif filename.endswith(('.docx', '.doc')):
                if not _DOCX_AVAILABLE:
                    return None, "Thư viện python-docx chưa được cài đặt. Chạy: pip install python-docx"
                docx_bytes = io.BytesIO(file.read())
                doc = python_docx.Document(docx_bytes)
                lines = []
                for para in doc.paragraphs:
                    stripped = para.text.strip()
                    if not stripped:
                        continue
                    style_name = para.style.name if para.style else ""
                    if style_name.startswith("Heading"):
                        try:
                            level = int(style_name.split()[-1])
                        except ValueError:
                            level = 1
                        lines.append("#" * level + " " + stripped)
                    else:
                        lines.append(stripped)
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_texts:
                            lines.append(" | ".join(row_texts))
                text = "\n".join(lines)

            else:
                return None, "Only PDF, DOCX, TXT, and MD files are supported."

            return text.strip(), None

        except Exception as e:
            return None, str(e)

    # ── Document Classification ───────────────────────────────────────────────

    @staticmethod
    def classify_document(text: str) -> dict:
        """
        Phân loại tài liệu bằng heuristic keyword matching.
        Returns { "type": "teaching_content" | "pedagogical_script", "reason": str }
        """
        text_normalized = text.lower()
        keywords = [
            "ngày soạn", "ngày dạy", "tiến trình dạy học", "thiết bị dạy học",
            "hoạt động khởi động", "hình thành kiến thức", "hoạt động luyện tập",
            "hoạt động vận dụng", "hoạt động của gv", "hoạt động của hs",
            "chuyển giao nhiệm vụ", "dự kiến sản phẩm", "giáo án", "lesson plan",
            "tiến trình bài dạy", "tổ chức thực hiện", "sản phẩm học tập",
            "bước 1:", "bước 2:", "bước 3:", "bước 4:",
            "mục tiêu:", "nội dung:", "tổ chức hoạt động"
        ]
        match_count = sum(1 for kw in keywords if kw in text_normalized)

        if match_count >= 2:
            return {
                "type": "pedagogical_script",
                "reason": f"Phát hiện nhanh {match_count} từ khóa đặc trưng của kịch bản sư phạm/giáo án.",
            }
        return {
            "type": "teaching_content",
            "reason": "Mặc định (Nội dung giảng dạy)",
        }

    # ── Helper: parse list field từ form ─────────────────────────────────────

    @staticmethod
    def parse_list_field(raw_val: str) -> list:
        if not raw_val:
            return []
        raw_val = raw_val.strip()
        if raw_val.startswith("[") and raw_val.endswith("]"):
            try:
                return json.loads(raw_val)
            except Exception:
                pass
        return [item.strip() for item in raw_val.split(",") if item.strip()]
