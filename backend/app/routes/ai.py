from flask_openapi3 import APIBlueprint, FileStorage
from flask import request, jsonify, current_app
from pydantic import BaseModel, Field
from typing import Optional
import pypdf
import io
try:
    import docx as python_docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from app.services.openrouter_service import OpenRouterService
from app.services.chunking_service import get_all_section_contexts, semantic_chunk, extract_outline

bp = APIBlueprint('ai', __name__)

# ─── Pydantic Models ──────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    prompt: Optional[str] = Field(None, description="A single user prompt string.")
    messages: Optional[list] = Field(None, description="Full message history (role/content pairs).")
    model: Optional[str] = Field(None, description="Optional OpenRouter model override.")
    temperature: Optional[float] = Field(0.7, description="Sampling temperature.")
    max_tokens: Optional[int] = Field(1000, description="Maximum tokens to generate.")

class MapRequest(BaseModel):
    extracted_knowledge: str = Field(..., description="Raw text extracted from source document (output of Feature 1).")
    system_template: str = Field(..., description="Lesson node template string (e.g. 'Node 1: Warm-up, Node 2: Core Theory').")
    model: Optional[str] = Field(None, description="Optional OpenRouter model override.")
    temperature: Optional[float] = Field(0.3, description="Sampling temperature (default 0.3).")

class EnrichRequest(BaseModel):
    mapped_nodes: object = Field(..., description="Mapped node array or JSON string (output of Feature 2).")
    rag_activities: Optional[object] = Field(None, description="List of teaching methodology strings from RAG context.")
    classroom_ctx: Optional[dict] = Field(None, description="Optional classroom context configuration.")
    model: Optional[str] = Field(None, description="Optional OpenRouter model override.")
    temperature: Optional[float] = Field(0.5, description="Sampling temperature (default 0.5).")

class ExtractRequest(BaseModel):
    file: FileStorage = Field(..., description="PDF, TXT, or MD file to extract text from")

class PipelineRequest(BaseModel):
    file: FileStorage = Field(..., description="PDF, TXT, or MD file to process through the full 3-phase pipeline")

class PipelineAdvancedRequest(BaseModel):
    file: FileStorage = Field(..., description="PDF, TXT, or MD file")
    use_chunking: Optional[bool] = Field(True, description="Dùng semantic chunking thay vì đưa full text vào prompt (mặc định: True)")
    k_chunks: Optional[int] = Field(5, description="Số chunks top-K lấy per section (mặc định: 5)")
    use_parallel: Optional[bool] = Field(True, description="Enrich các node song song thay vì tuần tự (mặc định: True)")
    enable_faithfulness: Optional[bool] = Field(False, description="Bật kiểm tra faithfulness/grounding sau khi sinh (mặc định: False, tốn thêm LLM call)")

# Default pedagogical template used by the pipeline
_DEFAULT_SYSTEM_TEMPLATE = (
    "Khung bài học 3 phần:\n"
    "Node 1: Khởi động (Warm-up) - Kích hoạt kiến thức nền và tạo hứng thú cho học sinh\n"
    "Node 2: Hình thành kiến thức (Core Theory) - Giới thiệu, giải thích và xây dựng các khái niệm, kiến thức mới cốt lõi\n"
    "Node 3: Luyện tập (Practice) - Học sinh thực hành, làm các bài tập củng cố ngay tại lớp"
)

# Default RAG activities pool used by the pipeline
_DEFAULT_RAG_ACTIVITIES = [
    "1. Trò chơi Kahoot câu hỏi trắc nghiệm ôn tập nhanh (5-7 phút)",
    "2. Thảo luận nhóm tranh biện (Think-Pair-Share)",
    "3. Sơ đồ tư duy tiếp sức theo nhóm (Mind Map Relay)",
    "4. Thẻ câu hỏi xoay vòng (Quiz Cards)",
    "5. Thí nghiệm hoặc mô phỏng thực hành có hướng dẫn",
]

# Default node names matching the system template
_DEFAULT_NODE_NAMES = [
    "Khởi động (Warm-up)",
    "Hình thành kiến thức (Core Theory)",
    "Luyện tập (Practice)",
]

# ─── Helper: extract text from uploaded file ─────────────────────────────────

def _extract_text_from_file(file) -> tuple[str | None, str | None]:
    """Returns (extracted_text, error_message)."""
    filename = file.filename.lower()
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    try:
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        if size > MAX_FILE_SIZE:
            return None, "Kích thước file vượt quá giới hạn cho phép (tối đa 50MB)."
    except Exception:
        pass

    try:
        if filename.endswith('.pdf'):
            pdf_bytes = io.BytesIO(file.read())
            reader = pypdf.PdfReader(pdf_bytes)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        elif filename.endswith(('.txt', '.md')):
            text = file.read().decode('utf-8', errors='ignore')
        elif filename.endswith(('.docx', '.doc')):
            if not _DOCX_AVAILABLE:
                return None, "Thư viện python-docx chưa được cài đặt. Chạy: pip install python-docx"
            docx_bytes = io.BytesIO(file.read())
            doc = python_docx.Document(docx_bytes)
            lines = []
            # Trích xuất paragraphs (giữ nguyên heading cấu trúc)
            for para in doc.paragraphs:
                stripped = para.text.strip()
                if not stripped:
                    continue
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except ValueError:
                        level = 1
                    lines.append("#" * level + " " + stripped)
                else:
                    lines.append(stripped)
            # Trích xuất nội dung bảng
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        lines.append(" | ".join(row_texts))
            text = "\n".join(lines)
        else:
            return None, "Only PDF, DOCX, TXT, and MD files are supported."
        return text.strip(), None
    except Exception as e:
        return None, str(e)

# ─── Routes ───────────────────────────────────────────────────────────────────

@bp.post(
    '/chat',
    summary="General Chat Completion",
    description="Send a prompt or message history to OpenRouter and receive an AI-generated response."
)
def chat(body: ChatRequest):
    """POST /api/ai/chat"""
    prompt = body.prompt
    messages = body.messages
    model = body.model
    temperature = body.temperature if body.temperature is not None else 0.7
    max_tokens = body.max_tokens if body.max_tokens is not None else 1000

    if not prompt and not messages:
        return jsonify({"error": "Bad Request", "message": "Either 'prompt' or 'messages' must be provided."}), 400
    if not messages:
        messages = [{"role": "user", "content": prompt}]
    elif not isinstance(messages, list):
        return jsonify({"error": "Bad Request", "message": "'messages' must be a JSON array."}), 400

    result = OpenRouterService.generate_chat_completion(messages=messages, model=model, temperature=temperature, max_tokens=max_tokens)
    if result.get("success"):
        return jsonify({"success": True, "response": result.get("content"), "model": result.get("model"), "usage": result.get("usage")}), 200
    else:
        error_msg = result.get("error", "Unknown error")
        return jsonify({"success": False, "error": error_msg}), 403 if "OPENROUTER_API_KEY is not set" in error_msg else 500


@bp.post(
    '/pedagogy/extract',
    summary="Feature 1: Extract Raw Text (No AI)",
    description=(
        "Upload a file (PDF, TXT, or MD) and extract its raw text content programmatically "
        "using `pypdf` or standard file reading. **No AI is used in this step.**"
    )
)
def extract_pedagogy(form: ExtractRequest):
    """POST /api/ai/pedagogy/extract"""
    file = form.file
    if not file or file.filename == '':
        return jsonify({"error": "Bad Request", "message": "No selected file."}), 400

    text, err = _extract_text_from_file(file)
    if err:
        status = 415 if "Only PDF" in err else 500
        return jsonify({"success": False, "error": err}), status

    return jsonify({"success": True, "filename": file.filename, "extracted_text": text}), 200


@bp.post(
    '/pedagogy/map',
    summary="Feature 2: Map Knowledge to Template (AI)",
    description=(
        "Map extracted knowledge points into the lesson nodes defined by a system template. "
        "Uses OpenRouter AI with a structured pedagogical prompt.\n\n"
        "**Input**: `extracted_knowledge` (text from Feature 1) + `system_template` (node definitions)\n\n"
        "**Output**: A JSON array of mapped nodes with `node_name`, `node_intent`, and `mapped_knowledge`."
    )
)
def map_pedagogy(body: MapRequest):
    """POST /api/ai/pedagogy/map"""
    extracted_knowledge = body.extracted_knowledge
    system_template = body.system_template
    model = body.model
    temperature = body.temperature if body.temperature is not None else 0.3

    if not extracted_knowledge or not system_template:
        return jsonify({"error": "Bad Request", "message": "Both 'extracted_knowledge' and 'system_template' are required."}), 400

    outline = extract_outline(extracted_knowledge)
    result = OpenRouterService.map_knowledge_to_template(
        extracted_knowledge=extracted_knowledge,
        system_template=system_template,
        model=model,
        temperature=temperature,
        outline=outline
    )
    if result.get("success"):
        return jsonify({"success": True, "mapped_nodes": result.get("content")}), 200
    return jsonify({"success": False, "error": result.get("error")}), 500


@bp.post(
    '/pedagogy/enrich',
    summary="Feature 3: Enrich Nodes with Activities (AI)",
    description=(
        "For each mapped lesson node, select and adapt the most suitable teaching methodology "
        "from the provided RAG activities. Uses OpenRouter AI.\n\n"
        "**Input**: `mapped_nodes` (output of Feature 2) + `rag_activities` (list of methods)\n\n"
        "**Output**: Final pedagogical script with `node_name`, `applied_activity`, and `execution_steps`."
    )
)
def enrich_pedagogy(body: EnrichRequest):
    """POST /api/ai/pedagogy/enrich"""
    mapped_nodes = body.mapped_nodes
    rag_activities = body.rag_activities
    classroom_ctx = body.classroom_ctx
    model = body.model
    temperature = body.temperature if body.temperature is not None else 0.5

    if not mapped_nodes:
        return jsonify({"error": "Bad Request", "message": "'mapped_nodes' is required."}), 400

    if not rag_activities and not classroom_ctx:
        return jsonify({"error": "Bad Request", "message": "Either 'rag_activities' or 'classroom_ctx' is required."}), 400

    result = OpenRouterService.enrich_nodes_with_activities(
        mapped_nodes=mapped_nodes,
        rag_activities=rag_activities,
        model=model,
        temperature=temperature,
        classroom_ctx=classroom_ctx
    )
    if result.get("success"):
        return jsonify({"success": True, "enriched_script": result.get("content")}), 200
    return jsonify({"success": False, "error": result.get("error")}), 500


@bp.post(
    '/pedagogy/pipeline',
    summary="Full Pipeline: Extract → Map → Enrich (Sync, Legacy)",
    description=(
        "**Automated 3-phase pipeline** — upload a file and receive the final enriched pedagogical script in one call.\n\n"
        "1. **Extract**: Reads raw text from the uploaded file (no AI)\n"
        "2. **Map**: AI maps extracted content into a built-in 3-node lesson template (Warm-up, Theory, Practice)\n"
        "3. **Enrich**: AI assigns interactive teaching activities to each node\n\n"
        "**Input**: `multipart/form-data` with a single `file` field (PDF, TXT, or MD).\n\n"
        "**Output**: `raw_extracted_text`, `mapped_nodes`, and `final_pedagogical_script` (Feature 3 output).\n\n"
        "⚡ **For real-time streaming + semantic chunking + parallel enrichment**, use `/pedagogy/stream` instead."
    )
)
def pipeline_pedagogy(form: PipelineRequest):
    """POST /api/ai/pedagogy/pipeline — Legacy sync pipeline (backward-compatible)"""
    file = form.file
    system_template = _DEFAULT_SYSTEM_TEMPLATE
    rag_activities = _DEFAULT_RAG_ACTIVITIES
    model = None

    if not file or file.filename == '':
        return jsonify({"error": "Bad Request", "message": "No selected file."}), 400

    # ── Phase 1: Extract (No AI) ──────────────────────────────────────────────
    extracted_text, err = _extract_text_from_file(file)
    if err:
        return jsonify({"success": False, "error": f"Extraction failed: {err}"}), 415 if "Only PDF" in err else 500
    if not extracted_text:
        return jsonify({"error": "Unprocessable Entity", "message": "Extracted text is empty."}), 422

    # ── Phase 1b: Semantic Chunking (mặc định bật) ───────────────────────────
    chunks = semantic_chunk(extracted_text)
    section_contexts = get_all_section_contexts(
        text=extracted_text,
        node_names=_DEFAULT_NODE_NAMES,
        k=5,
    )

    # ── Phase 2: Map (AI) với section-aware context ───────────────────────────
    outline = extract_outline(extracted_text)
    map_result = OpenRouterService.map_knowledge_to_template(
        extracted_knowledge=extracted_text,
        system_template=system_template,
        model=model,
        temperature=0.3,
        use_chunking=True,
        section_contexts=section_contexts,
        outline=outline
    )
    if not map_result.get("success"):
        return jsonify({"success": False, "error": f"Mapping phase failed: {map_result.get('error')}"}), 500

    mapped_nodes = map_result.get("content")

    # ── Phase 3: Enrich (AI) — backward-compat sequential ─────────────────────
    enrich_result = OpenRouterService.enrich_nodes_with_activities(
        mapped_nodes=mapped_nodes,
        rag_activities=rag_activities,
        model=model,
        temperature=0.5
    )
    if not enrich_result.get("success"):
        return jsonify({"success": False, "error": f"Enrichment phase failed: {enrich_result.get('error')}"}), 500

    return jsonify({
        "success": True,
        "raw_extracted_text": extracted_text,
        "chunks_info": {
            "total_chunks": len(chunks),
            "sections": list(section_contexts.keys()),
        },
        "mapped_nodes": mapped_nodes,
        "final_pedagogical_script": enrich_result.get("content")
    }), 200


@bp.get(
    '/config',
    summary="AI Config Status",
    description="Check whether the OpenRouter API key is configured and what the default model is."
)
def get_config():
    """GET /api/ai/config"""
    return jsonify({
        "api_key_configured": bool(current_app.config.get('OPENROUTER_API_KEY')),
        "default_model": current_app.config.get('OPENROUTER_MODEL')
    }), 200
